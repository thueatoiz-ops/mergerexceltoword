import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
import io
from pathlib import Path
import tempfile
import os
import re

def replace_placeholder_in_paragraph(paragraph, row_data):
    """Thay thế placeholder trong paragraph, giữ nguyên định dạng"""
    full_text = paragraph.text
    
    # Kiểm tra xem có placeholder nào không
    has_placeholder = False
    for key in row_data.keys():
        if f"{{{{{key}}}}}" in full_text:
            has_placeholder = True
            break
    
    if not has_placeholder:
        return
    
    # Thay thế trong toàn bộ text
    new_text = full_text
    for key, value in row_data.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in new_text:
            new_text = new_text.replace(
                placeholder,
                str(value) if pd.notna(value) else ""
            )
    
    # Cập nhật paragraph
    if new_text != full_text:
        # Xóa tất cả runs
        paragraph.clear()
        # Thêm run mới với text đã thay thế
        paragraph.add_run(new_text)

def replace_placeholder_in_table(table, row_data):
    """Thay thế placeholder trong table"""
    for row_table in table.rows:
        for cell in row_table.cells:
            # Xử lý paragraphs trong cell
            for paragraph in cell.paragraphs:
                replace_placeholder_in_paragraph(paragraph, row_data)

st.set_page_config(
    page_title="Trộn dữ liệu Excel sang Word",
    page_icon="📄",
    layout="wide"
)

st.title("📄 Tool Trộn Dữ Liệu Excel sang Word (Mail Merge)")
st.markdown("---")

# Sidebar hướng dẫn
with st.sidebar:
    st.header("📋 Hướng dẫn sử dụng")
    st.markdown("""
    1. **Upload file Excel** chứa dữ liệu cần trộn
    2. **Upload file Word template** (mẫu) với các placeholder như {{Tên}}, {{Tuổi}}, ...
    3. Click **Trộn dữ liệu** để tạo file Word
    4. Download file kết quả
    
    **Lưu ý:** 
    - Trong Word template, sử dụng cú pháp `{{TênCột}}` để đánh dấu vị trí cần điền dữ liệu
    - Ví dụ: `{{HọTên}}`, `{{ĐịaChỉ}}`, `{{SốĐiệnThoại}}`
    - Tất cả dòng dữ liệu sẽ được trộn vào một file Word duy nhất (giống Mail Merge)
    """)

# Upload files
col1, col2 = st.columns(2)

with col1:
    st.subheader("📊 File Excel (Dữ liệu nguồn)")
    excel_file = st.file_uploader(
        "Chọn file Excel",
        type=['xlsx', 'xls'],
        key="excel_upload"
    )

with col2:
    st.subheader("📝 File Word Template (Mẫu)")
    word_file = st.file_uploader(
        "Chọn file Word template",
        type=['docx'],
        key="word_upload"
    )

if excel_file and word_file:
    st.markdown("---")
    
    try:
        # Đọc Excel
        df = pd.read_excel(excel_file)
        
        # Hiển thị preview dữ liệu
        st.subheader("👀 Xem trước dữ liệu Excel")
        st.dataframe(df.head(10), use_container_width=True)
        st.info(f"Tổng số dòng dữ liệu: {len(df)}")
        
        # Cấu hình
        st.subheader("⚙️ Cấu hình trộn dữ liệu")
        
        col1, col2 = st.columns(2)
        
        with col1:
            # Hiển thị các cột có sẵn
            st.markdown("**Các cột trong Excel:**")
            st.write(list(df.columns))
            
            # Hiển thị ví dụ placeholder
            st.markdown("**Ví dụ placeholder trong Word:**")
            example_placeholders = [f"{{{{{col}}}}}" for col in df.columns[:3]]
            st.code("\n".join(example_placeholders))
        
        with col2:
            # Tùy chọn ngắt trang
            page_break = st.checkbox(
                "Thêm ngắt trang giữa các bản ghi",
                value=True,
                help="Nếu bật, mỗi bản ghi sẽ bắt đầu ở trang mới"
            )
            
            # Tùy chọn bỏ qua dòng trống
            skip_empty = st.checkbox(
                "Bỏ qua dòng có dữ liệu trống",
                value=False,
                help="Bỏ qua các dòng mà tất cả các cột đều trống"
            )
        
        # Nút trộn dữ liệu
        if st.button("🔄 Trộn dữ liệu (Mail Merge)", type="primary", use_container_width=True):
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            try:
                # Đọc Word template vào memory
                word_template_bytes = word_file.read()
                
                # Tạo document mới
                merged_doc = Document()
                processed_count = 0
                
                for idx, row in df.iterrows():
                    # Bỏ qua dòng trống nếu được chọn
                    if skip_empty and row.isna().all():
                        continue
                    
                    status_text.text(f"Đang xử lý dòng {idx + 1}/{len(df)}...")
                    progress_bar.progress((idx + 1) / len(df))
                    
                    # Tạo bản copy của template cho mỗi dòng
                    template_doc = Document(io.BytesIO(word_template_bytes))
                    
                    # Chuyển đổi row thành dict và xử lý NaN
                    row_data = {}
                    for key, value in row.items():
                        row_data[key] = value if pd.notna(value) else ""
                    
                    # Thay thế placeholder trong paragraphs
                    for paragraph in template_doc.paragraphs:
                        replace_placeholder_in_paragraph(paragraph, row_data)
                    
                    # Thay thế placeholder trong tables
                    for table in template_doc.tables:
                        replace_placeholder_in_table(table, row_data)
                    
                    # Thêm ngắt trang trước mỗi bản ghi (trừ bản ghi đầu tiên)
                    if processed_count > 0 and page_break:
                        # Thêm paragraph với page break
                        p = merged_doc.add_paragraph()
                        run = p.add_run()
                        run.add_break(WD_BREAK.PAGE)
                    
                    # Sao chép tất cả elements từ template vào document chính
                    for element in template_doc.element.body:
                        merged_doc.element.body.append(element)
                    
                    processed_count += 1
                
                # Lưu file
                output_buffer = io.BytesIO()
                merged_doc.save(output_buffer)
                output_buffer.seek(0)
                
                status_text.text("✅ Hoàn thành!")
                progress_bar.progress(1.0)
                
                st.success(f"✅ Đã trộn {processed_count} dòng dữ liệu vào một file Word thành công!")
                
                # Hiển thị thông tin file
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric("Số dòng đã trộn", processed_count)
                with col2:
                    st.metric("Kích thước file", f"{len(output_buffer.getvalue()) / 1024:.2f} KB")
                with col3:
                    st.metric("Chế độ", "Có ngắt trang" if page_break else "Nối tiếp")
                
                st.download_button(
                    label="📥 Download file Word đã trộn",
                    data=output_buffer,
                    file_name="merged_document.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )
            
            except Exception as e:
                st.error(f"❌ Lỗi khi trộn dữ liệu: {str(e)}")
                st.exception(e)
    
    except Exception as e:
        st.error(f"❌ Lỗi khi đọc file: {str(e)}")
        st.exception(e)

else:
    st.info("👆 Vui lòng upload cả file Excel và file Word template để bắt đầu")

# Footer
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: gray;'>
    <p>Tool Trộn Dữ Liệu Excel sang Word | Made with Streamlit</p>
</div>
""", unsafe_allow_html=True)

